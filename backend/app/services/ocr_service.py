import io
import os
import platform
import time
import uuid
from typing import List, Dict, Any, Tuple, Optional

import pytesseract
import pypdfium2 as pdfium
from PIL import Image, ImageOps, ImageFilter
from docx import Document
from pytesseract import Output

from app.models.schemas import OCRResponse, PageText
from app.services import file_service
from app.core.config import settings

from app.services.ocr_phase2_adapter import phase2_enrich_page
from app.services.semantic_cleanup import cleanup_page  # Phase 3 early (existing)
from app.services.quality_scoring import score_page

# Phase 3 proper
from app.services.document_normalizer import normalize_document

# Phase 4: optional multi-engine orchestration (docTR/TrOCR)
try:
    from app.services.engine_orchestrator import orchestrate_page_ocr
except Exception:
    orchestrate_page_ocr = None  # type: ignore

# Phase 4 exports regeneration (optional; depends on your repo files)
try:
    from app.services.export_markdown import document_to_markdown
except Exception:
    document_to_markdown = None  # type: ignore

try:
    from app.services.export_html import document_to_html
except Exception:
    document_to_html = None  # type: ignore

try:
    from app.services.chunking import chunk_document
except Exception:
    chunk_document = None  # type: ignore
# Diagnostics v2 (noise/skew/mixed-script)
try:
    from app.services.diagnostics_v2 import compute_page_diagnostics
except Exception:
    compute_page_diagnostics = None  # type: ignore



def configure_tesseract():
    env_cmd = os.getenv("TESSERACT_CMD")
    if env_cmd and os.path.exists(env_cmd):
        pytesseract.pytesseract.tesseract_cmd = env_cmd
        return

    if platform.system().lower().startswith("win"):
        win_default = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        if os.path.exists(win_default):
            pytesseract.pytesseract.tesseract_cmd = win_default
            return


configure_tesseract()


def _preprocess(image: Image.Image) -> Image.Image:
    image = image.convert("L")

    w, h = image.size
    max_side = max(w, h)
    if max_side < 1000:
        scale = 1000 / max_side
        image = image.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

    image = ImageOps.autocontrast(image)
    image = image.filter(ImageFilter.SHARPEN)
    return image


def ocr_image_words(image: Image.Image) -> Dict[str, Any]:
    """OCR with word-level confidence + bbox (NO dropping)."""
    image = _preprocess(image)
    data = pytesseract.image_to_data(image, output_type=Output.DICT)

    words: List[Dict[str, Any]] = []
    texts: List[str] = []

    n = len(data.get("text", []))
    for i in range(n):
        txt = (data["text"][i] or "").strip()
        conf = data["conf"][i]

        if not txt:
            continue

        try:
            conf_i = float(conf)
            conf_f = conf_i / 100.0 if conf_i >= 0 else None
        except Exception:
            conf_f = None

        left = int(data["left"][i])
        top = int(data["top"][i])
        width = int(data["width"][i])
        height = int(data["height"][i])

        words.append(
            {
                "text": txt,
                "confidence": conf_f,
                "bbox": [left, top, left + width, top + height],
            }
        )
        texts.append(txt)

    return {"text": " ".join(texts).strip(), "words": words}


def extract_from_pdf(file_bytes: bytes) -> Tuple[List[PageText], List[Optional[Image.Image]]]:
    """
    Returns:
      pages: List[PageText]
      page_images: aligned list (PIL image for image-rendered pages; None for text-extracted pages)
    """
    pdf = pdfium.PdfDocument(file_bytes)
    pages: List[PageText] = []
    page_images: List[Optional[Image.Image]] = []

    for i in range(len(pdf)):
        page = pdf[i]

        text = ""
        try:
            textpage = page.get_textpage()
            text = (textpage.get_text_range() or "").strip()
        except Exception:
            text = ""

        if text:
            pages.append(PageText(page_number=i + 1, text=text))
            page_images.append(None)
        else:
            scale = 300 / 72.0
            bitmap = page.render(scale=scale)
            pil_image = bitmap.to_pil().convert("RGB")

            o = ocr_image_words(pil_image)
            pages.append(PageText(page_number=i + 1, text=o["text"], words=o["words"]))
            page_images.append(pil_image)

    return pages, page_images


def extract_from_image(file_bytes: bytes) -> Tuple[List[PageText], List[Optional[Image.Image]]]:
    image = Image.open(io.BytesIO(file_bytes)).convert("RGB")
    o = ocr_image_words(image)
    return [PageText(page_number=1, text=o["text"], words=o["words"])], [image]


def extract_from_docx(file_bytes: bytes) -> Tuple[List[PageText], List[Optional[Image.Image]]]:
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    return [PageText(page_number=1, text=text)], [None]


def process_file(
    file_bytes: bytes,
    filename: str,
    document_type: str,
    *,
    zero_retention: bool | None = None,
) -> OCRResponse:
    start = time.time()
    job_id = str(uuid.uuid4())

    if zero_retention is None:
        zero_retention = settings.ZERO_RETENTION_DEFAULT

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    # --- Phase 1: ingestion (plus keep images for multi-engine) ---
    if ext == "pdf":
        pages, page_images = extract_from_pdf(file_bytes)
    elif ext in {"jpg", "jpeg", "png", "bmp", "tif", "tiff"}:
        pages, page_images = extract_from_image(file_bytes)
    elif ext == "docx":
        pages, page_images = extract_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext or 'unknown'}")

    enriched_pages: List[PageText] = []
    page_quality: List[Dict[str, Any]] = []

    # --- Phase 2 + Phase 3 early (non-destructive) ---
    for p in pages:
        page_dict = p.model_dump() if hasattr(p, "model_dump") else p.dict()

        # Phase 2: layout reconstruction
        page_dict = phase2_enrich_page(page_dict)

        # Phase 3 early: legacy cleanup + quality
        page_dict = cleanup_page(page_dict)
        page_dict["quality"] = score_page(
            page_dict.get("words") or [],
            page_dict.get("text_normalized") or page_dict.get("text") or "",
        )

        page_quality.append({"page": page_dict.get("page_number"), **(page_dict.get("quality") or {})})
        enriched_pages.append(PageText(**page_dict))

    full_text = "\n\n".join((p.text_normalized or p.text or "") for p in enriched_pages).strip()
    processing_time_ms = int((time.time() - start) * 1000)

    if not zero_retention:
        file_service.save_unique_by_name(filename, file_bytes)
    else:
        file_service.delete_if_exists(filename)

    qs = [q.get("quality_score") for q in page_quality if isinstance(q.get("quality_score"), (int, float))]
    avg_quality = (sum(qs) / len(qs)) if qs else None

    # --- Phase 3 proper: canonical document model (TOP-LEVEL) ---
    document_model = normalize_document(
        [p.model_dump() for p in enriched_pages],
        full_text=full_text,
    )

    # --- Diagnostics v2 (non-destructive) ---
    if compute_page_diagnostics is not None and isinstance(page_images, list) and document_model is not None:
        try:
            v2_pages = []
            for i, img in enumerate(page_images):
                if img is None:
                    continue
                page_num = i + 1
                # derive page text from enriched_pages if available
                page_text = ""
                try:
                    page_text = enriched_pages[i].text if i < len(enriched_pages) else ""
                except Exception:
                    page_text = ""
                v2_pages.append({"page_number": page_num, **compute_page_diagnostics(img, page_text)})
            # attach
            try:
                document_model.diagnostics.setdefault("v2", {})
                document_model.diagnostics["v2"]["pages"] = v2_pages
            except Exception:
                pass
        except Exception:
            pass

    # --- Phase 4: docTR + TrOCR orchestration (OPTIONAL, safe) ---
    # Runs only if (a) orchestrator exists and (b) we have page images
    if orchestrate_page_ocr is not None and isinstance(page_images, list) and document_model is not None:
        try:
            dm = document_model.model_dump() if hasattr(document_model, "model_dump") else dict(document_model)
            dm_pages = dm.get("pages") or []
            updated_pages = []
            for i, pg in enumerate(dm_pages):
                img = page_images[i] if i < len(page_images) else None
                if img is None:
                    updated_pages.append(pg)
                    continue
                page_number = int(pg.get("page_number") or (i + 1))
                updated_pages.append(
                    orchestrate_page_ocr(
                        page_image=img,
                        page_number=page_number,
                        base_page_dict=pg,
                    )
                )
            dm["pages"] = updated_pages

            # Rebuild full_text_normalized from blocks after overrides
            parts: List[str] = []
            for pg in dm["pages"]:
                for b in (pg.get("blocks") or []):
                    t = (b.get("text_normalized") or b.get("text") or "").strip()
                    if t:
                        parts.append(t)
            dm["full_text_normalized"] = "\n\n".join(parts).strip()

            # Regenerate exports if helpers exist
            if document_to_markdown is not None:
                dm["markdown"] = document_to_markdown(dm["pages"], tables=dm.get("tables") or [])
            if document_to_html is not None:
                dm["html"] = document_to_html(dm["pages"], tables=dm.get("tables") or [])
            if chunk_document is not None:
                dm["chunks"] = chunk_document(dm["pages"])

            # Try to reconstruct strongly-typed model
            try:
                from app.models.document_model import DocumentModel  # type: ignore
                document_model = DocumentModel(**dm)
            except Exception:
                document_model = dm  # type: ignore
        except Exception:
            # orchestration must never fail the request
            pass

    return OCRResponse(
        job_id=job_id,
        status="success",
        document_type=document_type,
        pages=enriched_pages,
        full_text=full_text,
        document=document_model,
        metadata={
            "file_name": filename,
            "file_type": ext,
            "num_pages": len(enriched_pages),
            "processing_time_ms": processing_time_ms,
            "engine": "tesseract(+doctr/trocr)",
            "zero_retention": bool(zero_retention),
            "phase2_complete": True,
            "phase3_complete": True,
            "avg_quality_score": avg_quality,
            "page_quality": page_quality,
        },
    )
